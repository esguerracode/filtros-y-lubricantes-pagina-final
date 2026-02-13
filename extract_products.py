#!/usr/bin/env python3
"""
Script para extraer productos del archivo KITS.docx
Genera un CSV listo para importar en WooCommerce
"""

import json
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import os
import re

def extract_products_from_docx(filepath):
    """
    Extrae productos del documento Word.
    Detecta tablas, imágenes y texto estructurado.
    """
    doc = Document(filepath)
    products = []
    
    print(f"\n{'='*60}")
    print(f"📄 PROCESANDO: {os.path.basename(filepath)}")
    print(f"{'='*60}\n")
    
    # Extraer todos los párrafos de texto
    all_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            all_text.append(para.text.strip())
    
    print("📝 TEXTO EXTRAÍDO:")
    print("\n".join(all_text[:50]))  # Primeras 50 líneas
    print("\n" + "="*60 + "\n")
    
    # Extraer tablas
    print(f"📊 TABLAS ENCONTRADAS: {len(doc.tables)}\n")
    for idx, table in enumerate(doc.tables):
        print(f"Tabla {idx + 1}:")
        for row_idx, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            print(f"  Fila {row_idx + 1}: {row_data}")
        print()
    
    # Extraer información de imágenes incrustadas
    print(f"🖼️ IMÁGENES EN EL DOCUMENTO:")
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
            print(f"  - Imagen {image_count}: {rel.target_ref}")
    
    print(f"\nTotal de imágenes: {image_count}")
    print("\n" + "="*60 + "\n")
    
    # Guardar todo el contenido en un archivo de texto para análisis
    with open('kits_content_analysis.txt', 'w', encoding='utf-8') as f:
        f.write("CONTENIDO COMPLETO DEL DOCUMENTO KITS.docx\n")
        f.write("=" * 80 + "\n\n")
        f.write("\n".join(all_text))
    
    print("✅ Contenido guardado en: kits_content_analysis.txt")
    
    return products

if __name__ == "__main__":
    docx_path = r"d:\00_CENTRAL\01_ACTIVO\Antigravity (No Borrar)\01_TRABAJO_ACTUAL\Filtros-Lubricantes-Web\Filtros-y-lubricantes-e-commerce\Filtros-lubricantes-pagina-final\KITS.docx"
    
    if not os.path.exists(docx_path):
        print(f"❌ ERROR: No se encontró el archivo {docx_path}")
        exit(1)
    
    products = extract_products_from_docx(docx_path)
